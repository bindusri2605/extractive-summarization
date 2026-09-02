import streamlit as st
import numpy as np
import pandas as pd
import re
import networkx as nx
import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer

# ============================================
# 1. SETUP & MODEL CACHING
# ============================================
st.set_page_config(page_title="HiLegalSum UI", layout="wide")

@st.cache_resource
def load_resources():
    model = SentenceTransformer("sentence-transformers/all-mpnet-base-v2")
    vectorizer = TfidfVectorizer(max_features=4000, stop_words="english")
    return model, vectorizer

model, vectorizer = load_resources()

LEGAL_KEYWORDS = ["court", "plaintiff", "defendant", "act", "law", "section", "clause", "judgment", "order", "appeal"]

# ============================================
# 2. CORE LOGIC & FILE EXTRACTION
# ============================================
def simple_sent_tokenize(text):
    text = text.replace("\n", " ")
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

def extract_text_from_file(uploaded_file):
    """Extracts text from PDF or DOCX files (used by the summarizer tab)."""
    text = ""
    if uploaded_file.type == "application/pdf":
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text

def run_hi_legal_sum(text, k, w_sem, w_pos, w_tfidf, lambda_param):
    sents = simple_sent_tokenize(text)
    if len(sents) <= k:
        return sents

    tfidf_matrix = vectorizer.fit_transform(sents).toarray()
    emb = model.encode(sents, convert_to_numpy=True)
    sims = cosine_similarity(emb)

    G = nx.from_numpy_array(sims)
    try:
        centrality = nx.eigenvector_centrality_numpy(G)
        centrality_scores = np.array([centrality[i] for i in range(len(sents))])
    except:
        centrality_scores = np.ones(len(sents))

    pos_scores = np.array([1/np.sqrt(i+1) for i in range(len(sents))])
    tfidf_scores = tfidf_matrix.sum(axis=1)
    keyword_scores = np.array([sum([sent.lower().count(w) for w in LEGAL_KEYWORDS]) for sent in sents])

    total_score = (w_sem * centrality_scores) + (w_pos * pos_scores) + (w_tfidf * tfidf_scores) + (0.5 * keyword_scores)

    candidates = list(range(len(sents)))
    final_idx = []
    while len(final_idx) < k and candidates:
        mmr_scores = []
        for idx in candidates:
            rel = total_score[idx]
            div = 0 if not final_idx else np.max(sims[idx][final_idx])
            mmr_scores.append(lambda_param * rel - (1 - lambda_param) * div)
        best_idx = candidates[np.argmax(mmr_scores)]
        final_idx.append(best_idx)
        candidates.remove(best_idx)

    return [sents[i] for i in sorted(final_idx)]


# ============================================
# 2b. LEGAL SUGGESTIONS LOGIC (fact-anchored analysis, NOT keyword triggering)
# ============================================
# No external LLM/API is available in this project. This engine deliberately
# does NOT fire just because a generic legal word ("court", "appeal",
# "breach", "some", ...) appears somewhere in the document. Each rule below
# requires a keyword AND a concrete, nearby anchor fact -- a day count, a
# date, a currency amount, a percentage, a named clause/section, a named
# law, a named jurisdiction, a specific ground, a specific duration, etc.
# Only when both the keyword and its anchor are found does a suggestion get
# generated, and the suggestion always names the specific fact it found so
# the reader can see exactly what triggered it. If a keyword appears with no
# such anchor nearby, nothing is generated for it -- that avoids generic,
# could-apply-to-any-document suggestions.

AMBIGUITY_MARKERS = [
    "tbd", "to be decided", "not sure", "unsure", "flexible",
    "as needed", "if possible", "might", "could be", "roughly",
    "reasonable efforts", "best efforts", "approximately"
]

DATE_PATTERN = re.compile(
    r'\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2}|'
    r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b',
    re.IGNORECASE
)
CURRENCY_PATTERN = re.compile(r'[\$₹€£]\s?\d[\d,]*(?:\.\d+)?')
PERCENT_PATTERN = re.compile(r'\b\d{1,3}(?:\.\d+)?\s?%')

_WIN = r'(?:(?![.!?]).){0,140}?'  # bounded "same sentence-ish" window, non-greedy

def _first_group(match):
    if not match:
        return None
    for g in match.groups():
        if g:
            return g.strip()
    return None


def analyze_legal_context(full_text):
    """
    Fact-anchored analysis of the ORIGINAL legal context. A suggestion is
    only produced when a keyword AND a concrete supporting fact are both
    found near each other in the text -- bare keyword presence alone never
    produces a suggestion. Returns dict: category -> list of insight strings
    (2-4 each, capped, de-duplicated), each naming the specific fact found.
    """
    low = full_text.lower()
    dates = list(dict.fromkeys(DATE_PATTERN.findall(full_text)))
    amounts = list(dict.fromkeys(CURRENCY_PATTERN.findall(full_text)))

    categories = {
        "Important Legal Issues": [],
        "Key Obligations": [],
        "Potential Risks": [],
        "Important Deadlines / Clauses": [],
        "Recommended Considerations": [],
    }
    # collected specific tokens, used to make Recommended Considerations
    # concrete rather than generic boilerplate
    deadline_specifics, risk_specifics, obligation_specifics = [], [], []

    # ---------- Important Legal Issues ----------
    # Litigation terms only matter here if tied to a specific clause reference.
    m = re.search(
        r'\b(court|plaintiff|defendant|appeal|judgment)\b' + _WIN + r'\b(?:Section|Article|Clause)\s+(\d+(?:\.\d+)*)\b'
        r'|\b(?:Section|Article|Clause)\s+(\d+(?:\.\d+)*)\b' + _WIN + r'\b(court|plaintiff|defendant|appeal|judgment)\b',
        full_text, re.IGNORECASE
    )
    if m:
        clause_num = m.group(2) or m.group(3)
        categories["Important Legal Issues"].append(
            f"Litigation-related language is tied to Section/Clause {clause_num} — confirm the procedural steps and deadlines set out there are being followed before any court action is pursued."
        )

    # Compliance only matters here if a specific named law/regulation is identifiable.
    named_law = re.search(
        r'\b([A-Z][A-Za-z]+(?:\s[A-Z][A-Za-z]+){0,4}\s(?:Act|Regulation|Code|Directive))\b|\b(GDPR|HIPAA|CCPA)\b',
        full_text
    )
    if named_law:
        law_name = named_law.group(1) or named_law.group(2)
        categories["Important Legal Issues"].append(
            f"The document is subject to {law_name} — confirm current compliance, since this specific law (not a generic compliance clause) governs how violations are treated."
        )

    if len(dates) > 1:
        categories["Important Legal Issues"].append(
            f"The document states multiple different dates ({', '.join(dates[:4])}) — confirm which one actually governs, since relying on the wrong one could misstate a deadline or effective date."
        )
    if len(amounts) > 1:
        categories["Important Legal Issues"].append(
            f"The document states multiple different monetary figures ({', '.join(amounts[:4])}) — confirm which figure is authoritative before relying on either."
        )

    conf_m = re.search(
        r'\bconfidential\w*\b' + _WIN + r'(\d{1,3}\s?(?:year|years|month|months))\b'
        r'|(\d{1,3}\s?(?:year|years|month|months))\b' + _WIN + r'\bconfidential\w*\b',
        full_text, re.IGNORECASE
    )
    if conf_m:
        dur = _first_group(conf_m)
        categories["Important Legal Issues"].append(
            f"Confidentiality obligations run for {dur} — confirm this duration and whether it survives after the agreement ends, since disclosure after that window may fall outside protection."
        )

    # ---------- Key Obligations ----------
    # Only include a modal-verb obligation when it has an identifiable, non-trivial action attached.
    obligation_hits = []
    for om in re.finditer(r'\b(?:shall|must)\s+((?:(?!\b(?:shall|must)\b)[^,.;]){3,90})', full_text, re.IGNORECASE):
        phrase = om.group(1).strip()
        phrase = re.sub(r'\s+', ' ', phrase)
        # Skip phrases that are really about a day-count deadline (notice/payment/etc.) --
        # those are already surfaced under "Important Deadlines / Clauses" and repeating
        # them here would duplicate the same point across two categories.
        if len(phrase.split()) >= 3 and not re.search(r'\b\d{1,3}\s*-?\s*days?\b', phrase, re.IGNORECASE):
            obligation_hits.append(phrase)
    seen_ob = []
    for p in obligation_hits:
        low_p = p.lower()
        if low_p not in seen_ob:
            seen_ob.append(low_p)
    for phrase in seen_ob[:3]:
        obligation_specifics.append(phrase)
        categories["Key Obligations"].append(
            f"The document imposes a mandatory obligation to {phrase} — confirm which party is responsible for this and whether it has actually been carried out."
        )

    assign_m = re.search(
        r'\bassign\w*\b' + _WIN + r'\b(written consent|prior consent|consent of|prior written approval)\b'
        r'|\b(written consent|prior consent|consent of|prior written approval)\b' + _WIN + r'\bassign\w*\b',
        full_text, re.IGNORECASE
    )
    if assign_m:
        req = _first_group(assign_m)
        categories["Key Obligations"].append(
            f"Assignment of rights or obligations requires {req} — confirm that requirement was actually obtained before treating any assignment as valid."
        )

    indem_m = re.search(
        r'\bindemnif\w*\b' + _WIN + r'\b(third[- ]part(?:y|ies)|losses|claims|attorney(?:s)?\'?\s?fees)\b',
        full_text, re.IGNORECASE
    )
    if indem_m:
        scope = indem_m.group(1)
        categories["Key Obligations"].append(
            f"An indemnification obligation specifically covers {scope} — identify which party bears this obligation and confirm the scope matches what actually occurred."
        )

    # ---------- Potential Risks ----------
    breach_m = re.search(
        r'\b(?:breach|default)\w*\b' + _WIN + r'\b(cure period|cure|remedy|remedies|\d{1,3}\s?days?\s?(?:to cure|notice))\b',
        full_text, re.IGNORECASE
    )
    if breach_m:
        detail = breach_m.group(1)
        risk_specifics.append(f"breach/default — {detail}")
        categories["Potential Risks"].append(
            f"A breach/default clause is tied to a specific '{detail}' provision — confirm the exact cure procedure and timeline before any remedy is pursued."
        )

    penalty_m = re.search(
        r'\b(?:penalty|penalties|damages|forfeit\w*)\b' + _WIN +
        r'([\$₹€£]\s?\d[\d,]*(?:\.\d+)?|\d{1,3}(?:\.\d+)?\s?%)'
        r'|([\$₹€£]\s?\d[\d,]*(?:\.\d+)?|\d{1,3}(?:\.\d+)?\s?%)' + _WIN + r'\b(?:penalty|penalties|damages|forfeit\w*)\b',
        full_text, re.IGNORECASE
    )
    if penalty_m:
        amt = _first_group(penalty_m)
        risk_specifics.append(f"penalty/damages of {amt}")
        categories["Potential Risks"].append(
            f"A penalty/damages provision specifies {amt} — assess this exact exposure against the value of the underlying obligation to judge how material the risk is."
        )

    liab_m = re.search(
        r'\b(?:liabl|liabilit)\w*\b' + _WIN +
        r'([\$₹€£]\s?\d[\d,]*(?:\.\d+)?|gross negligence|willful misconduct|cap(?:ped)?\s?at\s?[\$₹€£]?\s?\d[\d,]*)',
        full_text, re.IGNORECASE
    )
    if liab_m:
        detail = liab_m.group(1)
        risk_specifics.append(f"liability tied to {detail}")
        categories["Potential Risks"].append(
            f"Liability is specifically tied to {detail} — confirm whether this cap or carve-out would actually cover a realistic loss scenario under this agreement."
        )

    ground_m = re.search(
        r'\bterminat\w*\b' + _WIN + r'\b(for cause|for convenience|material breach|without cause)\b',
        full_text, re.IGNORECASE
    )
    if ground_m:
        ground = ground_m.group(1)
        risk_specifics.append(f"termination {ground}")
        categories["Potential Risks"].append(
            f"Termination is permitted '{ground}' — confirm the specific conditions for this ground were actually met, since invoking it without meeting them can itself be treated as a breach."
        )

    # ---------- Important Deadlines / Clauses ----------
    notice_m = re.search(
        r'(\d{1,3})\s*-?\s*(?:day|days)\b' + _WIN + r'\bterminat\w*\b'
        r'|\bterminat\w*\b' + _WIN + r'(\d{1,3})\s*-?\s*(?:day|days)\b',
        full_text, re.IGNORECASE
    )
    notice_days = _first_group(notice_m)
    if notice_days:
        deadline_specifics.append(f"{notice_days}-day notice-before-termination")
        categories["Important Deadlines / Clauses"].append(
            f"Termination requires a {notice_days}-day notice period first — verify that notice was actually given and that this exact window elapsed before termination was carried out."
        )

    payment_m = re.search(
        r'(\d{1,3})\s*-?\s*(?:day|days)\b' + _WIN + r'\bpa(?:y|yment)\w*\b'
        r'|\bpa(?:y|yment)\w*\b' + _WIN + r'(\d{1,3})\s*-?\s*(?:day|days)\b',
        full_text, re.IGNORECASE
    )
    payment_days = _first_group(payment_m)
    if payment_days:
        deadline_specifics.append(f"{payment_days}-day payment window")
        categories["Important Deadlines / Clauses"].append(
            f"Payment is due within {payment_days} days — confirm whether payment was made inside that exact window, since this specific clause (not a general delay) is what triggers late-payment consequences."
        )

    renew_m = re.search(
        r'\brenew\w*\b' + _WIN + r'(\d{1,3}\s*-?\s*(?:day|days))\b'
        r'|\brenew\w*\b' + _WIN + r'\b(successive|automatic\w*)\b',
        full_text, re.IGNORECASE
    )
    if renew_m:
        detail = _first_group(renew_m)
        deadline_specifics.append(f"renewal — {detail}")
        categories["Important Deadlines / Clauses"].append(
            f"The renewal clause specifically involves {detail} — confirm the exact deadline/method to opt out, since missing it would trigger this specific renewal mechanism."
        )

    fm_m = re.search(
        r'\bforce majeure\b' + _WIN + r'\b(natural disaster|pandemic|epidemic|war|strike|act of god|government action|labou?r dispute)\b',
        full_text, re.IGNORECASE
    )
    if fm_m:
        cause = fm_m.group(1)
        categories["Important Deadlines / Clauses"].append(
            f"The force majeure clause specifically lists '{cause}' as a covered event — confirm whether the actual circumstances at issue fall within this named list, since unlisted events may not be covered."
        )

    law_m = re.search(
        r'\b(?:governing law|laws of|jurisdiction of)\b' + _WIN + r'\b(?:State of\s+)?([A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+)?)\b'
        r'|\b(AAA|ICC|American Arbitration Association|International Chamber of Commerce)\b',
        full_text
    )
    if law_m:
        forum = law_m.group(1) or law_m.group(2)
        categories["Important Deadlines / Clauses"].append(
            f"The document specifies {forum} as the governing law/forum — confirm this is the forum actually being used, since disputes filed elsewhere may not be enforceable under this clause."
        )

    if dates and not notice_days and not payment_days:
        deadline_specifics.append(f"date {dates[0]}")
        categories["Important Deadlines / Clauses"].append(
            f"The document references the specific date {dates[0]} — confirm exactly what obligation or deadline attaches to it and calendar it accordingly."
        )

    # ---------- Recommended Considerations (built from what was actually found) ----------
    ambiguity_hits = [m for m in AMBIGUITY_MARKERS if re.search(rf'\b{re.escape(m)}\b', low)]
    if ambiguity_hits:
        categories["Recommended Considerations"].append(
            f"The document uses the undefined term '{ambiguity_hits[0]}' rather than a fixed standard — seek written clarification of what this specifically requires, since it's open to differing interpretation."
        )
    if deadline_specifics:
        categories["Recommended Considerations"].append(
            f"Calendar the specific deadlines identified in this document ({'; '.join(deadline_specifics[:3])}), since each is tied to a distinct consequence if missed."
        )
    if risk_specifics:
        categories["Recommended Considerations"].append(
            f"Given the specific risk provisions identified ({'; '.join(risk_specifics[:3])}), consider having a qualified attorney assess actual exposure under these exact terms before taking further action."
        )
    if obligation_specifics:
        categories["Recommended Considerations"].append(
            f"Retain documentation showing performance of the specific obligations identified ({'; '.join(obligation_specifics[:3])}), in case compliance with these exact terms is later disputed."
        )

    # de-duplicate and cap each category to 2-4 items
    for cat in categories:
        seen, unique = set(), []
        for item in categories[cat]:
            if item not in seen:
                seen.add(item)
                unique.append(item)
        categories[cat] = unique[:4]

    return categories


def render_legal_suggestions(full_text):
    """Renders the '💡 Legal Suggestions' section. Only fact-anchored,
    document-specific insights are shown; a category with no anchored
    finding says so explicitly rather than filling in generic advice."""
    st.markdown("### 💡 Legal Suggestions")

    categories = analyze_legal_context(full_text)

    if not any(categories.values()):
        st.write(
            "No suggestions with sufficient specific grounding (dates, amounts, "
            "named clauses, defined terms, etc.) were found in the provided text."
        )
    else:
        for cat_name, items in categories.items():
            st.markdown(f"**{cat_name}**")
            if items:
                for it in items:
                    st.markdown(f"- {it}")
            else:
                st.markdown("- No specific, fact-anchored issue was identified in this category from the provided context.")

    st.divider()
    st.caption(
        "These suggestions are generated from the provided document for "
        "informational purposes only and do not constitute legal advice."
    )


# ============================================
# 2c. REQUIREMENT ANALYSIS LOGIC (unchanged)
# ============================================
# No external LLM/API is available in this project, so the analysis below is
# built entirely from the libraries already used elsewhere in the app
# (regex, TF-IDF, sentence-transformer embeddings, simple keyword rules).

MAX_CHARS_PER_DOC = 50_000       # guard against extremely large documents
MAX_SENTENCES_FOR_SCAN = 600     # cap sentence-level scanning for performance

# Requirement "checklist" categories used to detect what's covered vs. missing.
REQUIREMENT_CATEGORIES = {
    "Objective / Goal": {
        "keywords": ["goal", "objective", "purpose", "aim", "why do we", "intended to"],
        "question": "What is the primary goal or business objective behind this requirement?",
        "suggestion": "State the objective in one sentence so every stakeholder can validate the solution against it."
    },
    "Scope": {
        "keywords": ["scope", "in scope", "out of scope", "boundary", "does not include", "excludes"],
        "question": "What is explicitly in scope, and what is out of scope?",
        "suggestion": "Add an explicit 'out of scope' list to prevent scope creep later."
    },
    "Users / Stakeholders": {
        "keywords": ["user", "audience", "stakeholder", "customer", "client", "end user", "persona"],
        "question": "Who are the primary users or stakeholders this requirement is for?",
        "suggestion": "Document the target user roles so the UI/UX and permissions can be tailored to them."
    },
    "Timeline": {
        "keywords": ["deadline", "timeline", "schedule", "due date", "by when", "go-live", "milestone"],
        "question": "What is the expected timeline, milestones, or deadline?",
        "suggestion": "Confirm a target date/milestones so effort can be planned realistically."
    },
    "Budget / Cost": {
        "keywords": ["budget", "cost", "pricing", "funding", "$", "₹", "€", "£"],
        "question": "Is there a budget or cost constraint that should shape the approach?",
        "suggestion": "Clarify budget constraints early to avoid proposing an over-scoped solution."
    },
    "Technical Constraints": {
        "keywords": ["technology", "platform", "stack", "framework", "database", "api", "integrat", "cloud", "on-premise"],
        "question": "Are there mandatory technologies, platforms, or systems this must integrate with?",
        "suggestion": "List any mandatory tech stack/integration constraints so the design fits existing infrastructure."
    },
    "Deliverables": {
        "keywords": ["deliverable", "output", "report", "artifact", "document to be produced", "final product"],
        "question": "What concrete deliverables or outputs are expected at the end?",
        "suggestion": "Define the exact deliverables (format, quantity) to make 'done' unambiguous."
    },
    "Success Criteria": {
        "keywords": ["success", "criteria", "kpi", "metric", "acceptance criteria", "measure"],
        "question": "How will success or completion be measured?",
        "suggestion": "Add measurable success/acceptance criteria to make the requirement testable."
    },
    "Compliance / Legal": {
        "keywords": ["compliance", "regulation", "legal", "policy", "gdpr", "security", "privacy", "confidential"],
        "question": "Are there compliance, legal, security, or privacy requirements to follow?",
        "suggestion": "Flag any regulatory/security requirements up front, since they usually affect architecture."
    },
    "Data Sources / Inputs": {
        "keywords": ["data source", "dataset", "input data", "existing data", "data feed", "database of"],
        "question": "What data sources or inputs will feed into this requirement?",
        "suggestion": "Identify the data sources now so access/permissions can be arranged early."
    },
}


def extract_text_generic(uploaded_file):
    """Extracts text from PDF, DOCX or TXT. Returns (text, error_message)."""
    try:
        name = uploaded_file.name.lower()
        if uploaded_file.type == "application/pdf" or name.endswith(".pdf"):
            text = ""
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text, None
        elif (uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              or name.endswith(".docx")):
            doc = Document(uploaded_file)
            text = "\n".join(p.text for p in doc.paragraphs)
            return text, None
        elif uploaded_file.type == "text/plain" or name.endswith(".txt"):
            raw = uploaded_file.read()
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                text = raw.decode("latin-1", errors="ignore")
            return text, None
        else:
            return "", f"Unsupported file type: {uploaded_file.name}"
    except Exception as e:
        return "", f"Could not read '{uploaded_file.name}' (file may be corrupted): {e}"


def clean_text(text):
    text = text.replace("\x00", " ")
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def chunk_text(text, max_chars=MAX_CHARS_PER_DOC):
    """Splits very large text into paragraph-respecting chunks and returns
    a single (possibly truncated) working copy plus a flag if truncation happened."""
    if len(text) <= max_chars:
        return text, False
    # Keep the beginning and end of the document rather than blindly cutting,
    # since key requirement info often lives at the start or end.
    head = text[: int(max_chars * 0.7)]
    tail = text[-int(max_chars * 0.3):]
    return head + "\n...\n[content truncated for length]\n...\n" + tail, True


def find_ambiguous_sentences(sentences):
    found = []
    for s in sentences:
        low = s.lower()
        for marker in AMBIGUITY_MARKERS:
            if re.search(rf'\b{re.escape(marker)}\b', low):
                found.append(s)
                break
    # de-duplicate while preserving order, cap to keep output concise
    seen = set()
    unique = []
    for s in found:
        if s not in seen:
            seen.add(s)
            unique.append(s)
    return unique[:8]


def find_contradictions(full_text):
    conflicts = []
    dates = list(dict.fromkeys(DATE_PATTERN.findall(full_text)))
    if len(dates) > 1:
        conflicts.append(f"Multiple different dates were mentioned ({', '.join(dates[:6])}). Please confirm which one applies.")
    amounts = list(dict.fromkeys(CURRENCY_PATTERN.findall(full_text)))
    if len(amounts) > 1:
        conflicts.append(f"Multiple different amounts were mentioned ({', '.join(amounts[:6])}). Please confirm the correct figure.")
    return conflicts


def analyze_requirement(description, doc_texts):
    """
    description: str
    doc_texts: dict[filename -> extracted text]
    Returns a dict with all sections needed for the UI, or raises on hard failure.
    """
    description = (description or "").strip()
    doc_texts = {k: v for k, v in doc_texts.items() if v and v.strip()}

    truncation_notes = []
    cleaned_docs = {}
    for fname, text in doc_texts.items():
        text = clean_text(text)
        text, truncated = chunk_text(text)
        if truncated:
            truncation_notes.append(fname)
        cleaned_docs[fname] = text

    combined_docs_text = "\n\n".join(cleaned_docs.values())
    full_text = (description + "\n\n" + combined_docs_text).strip()

    if not full_text:
        raise ValueError("No usable text found in the description or the uploaded documents.")

    desc_lower = description.lower()
    docs_lower = combined_docs_text.lower()

    info_identified = []
    missing_info = []
    clarifying_questions = []
    suggestions = []

    for category, cfg in REQUIREMENT_CATEGORIES.items():
        in_desc = any(kw in desc_lower for kw in cfg["keywords"])
        in_docs = any(kw in docs_lower for kw in cfg["keywords"])
        if in_desc or in_docs:
            source = []
            if in_desc:
                source.append("description")
            if in_docs:
                source.append("documents")
            info_identified.append((category, " & ".join(source)))
            suggestions.append(cfg["suggestion"])
        else:
            missing_info.append(category)
            clarifying_questions.append(cfg["question"])

    # Sentence-level scan for ambiguity / contradictions (capped for performance)
    all_sentences = simple_sent_tokenize(full_text)[:MAX_SENTENCES_FOR_SCAN]
    ambiguous = find_ambiguous_sentences(all_sentences)
    contradictions = find_contradictions(full_text)

    # Short extractive summary of the requirement, reusing the existing
    # HiLegalSum summarization engine already in this app.
    try:
        summary_k = 3 if len(all_sentences) > 3 else len(all_sentences)
        requirement_summary = run_hi_legal_sum(
            full_text, k=max(summary_k, 1), w_sem=1.0, w_pos=0.15, w_tfidf=0.25, lambda_param=0.7
        )
    except Exception:
        # Fall back to the first couple of sentences if the model call fails.
        requirement_summary = all_sentences[:3]

    return {
        "summary": requirement_summary,
        "info_identified": info_identified,
        "missing_info": missing_info,
        "clarifying_questions": clarifying_questions[:8],   # keep it concise & prioritized
        "suggestions": suggestions[:8],
        "ambiguous": ambiguous,
        "contradictions": contradictions,
        "truncation_notes": truncation_notes,
    }


# ============================================
# 3. STREAMLIT UI LAYOUT
# ============================================
st.title("🧾 HiLegalSum")
st.subheader("Legal-Aware Extractive Summarization & Requirement Analysis")

st.link_button(
    "🔗 View Source Code",
    "https://github.com/bindusri2605/extractive-summarization"
)

# Sidebar (used by the Summarization tab)
with st.sidebar:
    st.header("⚙️ Algorithm Settings")
    k_val = st.slider("Sentences in Summary", 1, 15, 5)
    st.markdown("---")
    w_sem = st.slider("Semantic Weight ($w_{sem}$)", 0.0, 2.0, 1.0)
    w_pos = st.slider("Position Weight ($w_{pos}$)", 0.0, 1.0, 0.15)
    w_tfidf = st.slider("TF-IDF Weight ($w_{tfidf}$)", 0.0, 1.0, 0.25)
    lambda_p = st.slider(r"MMR Diversity ($\lambda$)", 0.0, 1.0, 0.7)

tab_summary, tab_requirements = st.tabs(["📝 Legal Summarization", "🧭 Requirement Analysis"])

# ---------------- TAB 1: Legal Summarization ----------------
with tab_summary:
    st.markdown("### 📝 Input Legal Document")
    col_text, col_file = st.columns([4, 1])

    with col_file:
        uploaded_file = st.file_uploader("Add File", type=["pdf", "docx"], key="summary_file")

    initial_text = ""
    if uploaded_file is not None:
        initial_text = extract_text_from_file(uploaded_file)
        st.success("File loaded!")

    with col_text:
        input_text = st.text_area(
            "Paste your legal document or bill text here:",
            value=initial_text,
            height=400,
            placeholder="Example: SECTION 1. SHORT TITLE. This Act may be cited as...",
            key="summary_text"
        )

    if st.button("Generate Legal Summary", type="primary"):
        if not input_text.strip():
            st.error("Please provide some text to summarize.")
        else:
            with st.spinner("Processing legal nodes and graph centrality..."):
                summary = run_hi_legal_sum(input_text, k_val, w_sem, w_pos, w_tfidf, lambda_p)

            # 1) HiLegalSum Output -- existing extractive summarization algorithm, unchanged
            st.success("Summary Generated!")
            st.markdown("### 🔷 HiLegalSum Output")
            for sent in summary:
                st.markdown(f"**•** {sent}")

            # 2) Original / Summary statistics -- existing, unchanged
            st.divider()
            col1, col2 = st.columns(2)
            col1.metric("Original Sentences", len(simple_sent_tokenize(input_text)))
            col2.metric("Summary Sentences", len(summary))

            # 3) Legal Suggestions -- analysis of the ORIGINAL input_text, always after the summary/stats
            st.divider()
            render_legal_suggestions(input_text)

# ---------------- TAB 2: Requirement Analysis (unchanged) ----------------
with tab_requirements:
    st.markdown("### 📝 Describe Your Requirement")
    req_description = st.text_area(
        "Enter a clear description of what you want:",
        height=180,
        placeholder="Example: We need a portal that lets clients upload contracts and get an automated risk summary...",
        key="req_description"
    )

    st.markdown("### 📎 Upload Supporting Documents (optional)")
    req_files = st.file_uploader(
        "Upload one or more PDF, DOCX, or TXT files",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="req_files"
    )

    if req_files:
        st.caption(f"{len(req_files)} file(s) selected: " + ", ".join(f.name for f in req_files))

    analyze_clicked = st.button("Analyze Requirement", type="primary", key="analyze_btn")

    if analyze_clicked:
        # ---- Error handling: invalid / missing input ----
        if not req_description.strip() and not req_files:
            st.error("Please enter a description and/or upload at least one supporting document.")
        else:
            doc_texts = {}
            extraction_warnings = []

            if req_files:
                with st.spinner("Extracting text from uploaded documents..."):
                    for f in req_files:
                        try:
                            text, err = extract_text_generic(f)
                        except Exception as e:
                            text, err = "", f"Unexpected error reading '{f.name}': {e}"

                        if err:
                            extraction_warnings.append(err)
                        elif not text.strip():
                            extraction_warnings.append(f"'{f.name}' appears to be empty or the text could not be extracted (possibly a scanned/corrupted file).")
                        else:
                            doc_texts[f.name] = text

            for w in extraction_warnings:
                st.warning(w)

            if not req_description.strip() and not doc_texts:
                st.error("No usable text was found. Please check your description or the uploaded file(s).")
            else:
                try:
                    with st.spinner("Analyzing requirement against supporting documents..."):
                        result = analyze_requirement(req_description, doc_texts)
                except Exception as e:
                    st.error(f"Analysis failed due to an internal error: {e}")
                    result = None

                if result:
                    for fname in result["truncation_notes"]:
                        st.info(f"'{fname}' was very large, so only the beginning and end were analyzed.")

                    st.success("Analysis complete!")
                    st.divider()

                    st.markdown("### 📌 Requirement Summary")
                    if result["summary"]:
                        for s in result["summary"]:
                            st.markdown(f"**•** {s}")
                    else:
                        st.write("Not enough text to summarize.")

                    col_a, col_b = st.columns(2)

                    with col_a:
                        st.markdown("### ✅ Information Identified")
                        if result["info_identified"]:
                            for category, source in result["info_identified"]:
                                st.markdown(f"- **{category}** — found in *{source}*")
                        else:
                            st.write("No standard requirement details were confidently identified yet.")

                    with col_b:
                        st.markdown("### ❓ Missing Information")
                        if result["missing_info"]:
                            for category in result["missing_info"]:
                                st.markdown(f"- {category}")
                        else:
                            st.write("No major gaps detected against the standard checklist. 🎉")

                    st.markdown("### 🧩 Clarifying Questions")
                    if result["clarifying_questions"]:
                        for i, q in enumerate(result["clarifying_questions"], 1):
                            st.markdown(f"{i}. {q}")
                    else:
                        st.write("No clarifying questions — the requirement looks reasonably complete.")

                    st.markdown("### 💡 Suggestions")
                    if result["suggestions"]:
                        for s in result["suggestions"]:
                            st.markdown(f"- {s}")
                    else:
                        st.write("No additional suggestions at this time.")

                    st.markdown("### ⚠️ Assumptions / Conflicts")
                    if result["ambiguous"] or result["contradictions"]:
                        for c in result["contradictions"]:
                            st.markdown(f"- 🔴 **Possible contradiction:** {c}")
                        for a in result["ambiguous"]:
                            st.markdown(f"- 🟡 **Ambiguous statement:** \"{a}\"")
                    else:
                        st.write("No ambiguous statements or contradictions detected.")
